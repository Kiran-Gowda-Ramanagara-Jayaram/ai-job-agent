# utils/docx_resume.py
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

def _add_heading(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

def _add_bullets_md(doc, md_text: str):
    # Expect '- ' bullets in markdown
    for line in md_text.splitlines():
        line = line.strip()
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])

def build_ats_docx(company: str, title: str, resume_bullets_md: str,
                   base_resume_md: str, out_dir: str) -> str:
    """
    Create a simple ATS-friendly .docx:
    - Header: Target Role @ Company
    - Section: Tailored Highlights (from generated bullets)
    - Section: Core Skills (pulled from base resume headings if present)
    """
    Path(out_dir).mkdir(parents=True, exist_ok=True)
    doc = Document()

    _add_heading(doc, f"{title} @ {company}", size=18)

    _add_heading(doc, "Tailored Highlights", size=14)
    _add_bullets_md(doc, resume_bullets_md)

    # Optional: pull a simple Skills section from base resume lines starting with 'Skills:'
    skills = []
    for line in base_resume_md.splitlines():
        if line.lower().startswith("skills:"):
            skills.append(line.split(":", 1)[-1].strip())
    if skills:
        _add_heading(doc, "Core Skills", size=14)
        for s in skills[:2]:
            p = doc.add_paragraph()
            p.add_run(s)

    path = Path(out_dir) / "ATS_resume.docx"
    doc.save(str(path))
    return str(path)
